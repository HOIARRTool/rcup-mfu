# streamlit_app.py
# -*- coding: utf-8 -*-

import os
import re
import json
import html
from io import BytesIO
from datetime import datetime, date, time
from typing import Any, Dict, List, Optional, Tuple

import pandas as pd
import requests
import streamlit as st
import streamlit.components.v1 as components
import gspread

from docx import Document
from docx.shared import Inches
from google.oauth2.service_account import Credentials
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseUpload


# =========================
# CONFIG / CONSTANTS
# =========================

SHEET_COLUMNS = [
    "record_id",
    "unit_name",
    "app_title",
    "event_date",                # YYYY-MM-DD
    "event_time",                # HH:MM
    "process_step",              # สั่งใช้ยา / จัด/จ่ายยา / ให้ยา / ผู้ป่วยใช้ยาผิดวิธี
    "drug_name",
    "severity_level",            # A-I
    "incident_detail",
    "timeline_text",
    "initial_correction",
    "rca_text",
    "rca_image_filename",        # ชื่อไฟล์ภาพ
    "rca_image_drive_url",       # ลิงก์ไฟล์ภาพบน Google Drive
    "development_plan",
    "created_at",
    "created_by",
]

PROCESS_OPTIONS = ["สั่งใช้ยา", "จัด/จ่ายยา", "ให้ยา", "ผู้ป่วยใช้ยาผิดวิธี"]
SEVERITY_OPTIONS = list("ABCDEFGHI")


# =========================
# PAGE SETUP
# =========================

st.set_page_config(
    page_title="PHOIR",
    page_icon="🏡",
    layout="wide",
)


# =========================
# HELPER: READ CONFIG (ENV ONLY for Render)
# =========================

def _get_env(
    key: str,
    default: Optional[str] = None,
    aliases: Optional[List[str]] = None,
) -> Optional[str]:
    """ดึงค่าจาก Environment Variables เท่านั้น"""
    keys = [key] + (aliases or [])
    for k in keys:
        v = os.getenv(k)
        if v is not None and str(v).strip() != "":
            return str(v).strip()
    return default


def get_app_config() -> Dict[str, Any]:
    app_title = _get_env("APP_TITLE", "PHOIR_DEMO")
    unit_name = _get_env("UNIT_NAME", "unknown-unit")
    login_user = _get_env("APP_LOGIN_USERNAME", "")
    login_pass = _get_env("APP_LOGIN_PASSWORD", "")

    gsheet_url = _get_env("GSHEET_URL", "")
    worksheet_name = _get_env("GSHEET_WORKSHEET", "PHOIR_DEMO", aliases=["GHEET_WORKSHEET"])

    gcp_sa_json = _get_env("GCP_SERVICE_ACCOUNT_JSON", "", aliases=["GSHEET_CREDENTIALS_JSON"])
    gemini_api_key = _get_env("GEMINI_API_KEY", "")
    gdrive_folder_id = _get_env("GDRIVE_FOLDER_ID", "")

    return {
        "APP_TITLE": app_title,
        "UNIT_NAME": unit_name,
        "APP_LOGIN_USERNAME": login_user,
        "APP_LOGIN_PASSWORD": login_pass,
        "GSHEET_URL": gsheet_url,
        "GSHEET_WORKSHEET": worksheet_name,
        "GCP_SERVICE_ACCOUNT_JSON": gcp_sa_json,
        "GEMINI_API_KEY": gemini_api_key,
        "GDRIVE_FOLDER_ID": gdrive_folder_id,
    }


CFG = get_app_config()


# =========================
# STYLING
# =========================

st.markdown(
    """
<style>
.block-container { padding-top: 1.2rem; }
.small-muted { color: #6b7280; font-size: 0.88rem; }
.card {
    border: 1px solid #e5e7eb;
    border-radius: 14px;
    padding: 14px;
    background: #ffffff;
}
.section-title {
    font-size: 1.05rem;
    font-weight: 700;
    margin-bottom: .5rem;
}
</style>
    """,
    unsafe_allow_html=True,
)


# =========================
# LOGIN
# =========================

def ensure_auth_state():
    if "authenticated" not in st.session_state:
        st.session_state.authenticated = False
    if "login_username" not in st.session_state:
        st.session_state.login_username = ""
    if "show_fishbone_preview" not in st.session_state:
        st.session_state.show_fishbone_preview = False


def render_login():
    ensure_auth_state()

    st.markdown(f"# 🏡 {CFG['APP_TITLE']}")
    st.markdown(
        "<div class='small-muted'>บันทึกอุบัติการณ์ในสถานพยาบาลปฐมภูมิ</div>",
        unsafe_allow_html=True,
    )
    st.markdown("---")

    c1, c2, c3 = st.columns([1, 1.6, 1])
    with c2:
        st.markdown("## 🔐 เข้าสู่ระบบ")
        st.caption(f"หน่วยงาน: **{CFG['UNIT_NAME']}**")

        username = st.text_input("ชื่อผู้ใช้", key="login_user_input")
        password = st.text_input("รหัสผ่าน", type="password", key="login_pass_input")

        if st.button("เข้าสู่ระบบ", use_container_width=True):
            expected_user = CFG["APP_LOGIN_USERNAME"]
            expected_pass = CFG["APP_LOGIN_PASSWORD"]

            # ถ้ายังไม่ได้ตั้งค่า login ใน env ให้ bypass แบบ dev
            if not expected_user or not expected_pass:
                st.session_state.authenticated = True
                st.session_state.login_username = username or "dev-user"
                st.warning("ยังไม่ได้ตั้งค่า APP_LOGIN_USERNAME / APP_LOGIN_PASSWORD ใน ENV → เข้าแบบ dev mode")
                st.rerun()

            if username == expected_user and password == expected_pass:
                st.session_state.authenticated = True
                st.session_state.login_username = username
                st.success("เข้าสู่ระบบสำเร็จ ✅")
                st.rerun()
            else:
                st.error("ชื่อผู้ใช้หรือรหัสผ่านไม่ถูกต้อง")


# =========================
# GOOGLE API (Sheets + Drive)
# =========================

@st.cache_resource(show_spinner=False)
def get_google_credentials():
    sa_json_str = CFG["GCP_SERVICE_ACCOUNT_JSON"]
    if not sa_json_str:
        raise ValueError("ไม่พบ GCP_SERVICE_ACCOUNT_JSON ใน Environment Variables")

    try:
        creds_dict = json.loads(sa_json_str)
    except json.JSONDecodeError as e:
        raise ValueError(f"GCP_SERVICE_ACCOUNT_JSON ไม่ใช่ JSON ที่ถูกต้อง: {e}")

    scopes = [
        "https://www.googleapis.com/auth/spreadsheets",
        "https://www.googleapis.com/auth/drive",
    ]
    creds = Credentials.from_service_account_info(creds_dict, scopes=scopes)
    return creds


@st.cache_resource(show_spinner=False)
def get_gspread_client():
    creds = get_google_credentials()
    client = gspread.authorize(creds)
    return client


@st.cache_resource(show_spinner=False)
def get_drive_service():
    creds = get_google_credentials()
    return build("drive", "v3", credentials=creds, cache_discovery=False)


# =========================
# GOOGLE SHEETS
# =========================

@st.cache_resource(show_spinner=False)
def get_worksheet():
    gsheet_url = CFG["GSHEET_URL"]
    worksheet_name = CFG["GSHEET_WORKSHEET"]

    if not gsheet_url:
        raise ValueError("ไม่พบ GSHEET_URL ใน Environment Variables")

    client = get_gspread_client()
    sh = client.open_by_url(gsheet_url)

    try:
        ws = sh.worksheet(worksheet_name)
    except gspread.WorksheetNotFound:
        ws = sh.add_worksheet(title=worksheet_name, rows=1000, cols=60)

    # ensure header row
    header = ws.row_values(1)
    if not header:
        ws.append_row(SHEET_COLUMNS, value_input_option="USER_ENTERED")
    else:
        # ถ้าหัวตารางยังไม่ครบ ให้เติมเฉพาะคอลัมน์ที่ขาดแบบปลอดภัย
        missing_cols = [c for c in SHEET_COLUMNS if c not in header]
        if missing_cols:
            all_vals = ws.get_all_values()
            if all_vals:
                df_old = pd.DataFrame(all_vals[1:], columns=all_vals[0])
            else:
                df_old = pd.DataFrame(columns=[])

            for col in SHEET_COLUMNS:
                if col not in df_old.columns:
                    df_old[col] = ""

            # เก็บเฉพาะคอลัมน์ตามระบบปัจจุบัน
            df_old = df_old[SHEET_COLUMNS]

            ws.clear()
            ws.append_row(SHEET_COLUMNS, value_input_option="USER_ENTERED")
            if not df_old.empty:
                ws.append_rows(
                    df_old.fillna("").astype(str).values.tolist(),
                    value_input_option="USER_ENTERED",
                )

    return ws


def append_record_to_sheet(record: Dict[str, Any]) -> None:
    ws = get_worksheet()

    row = []
    for col in SHEET_COLUMNS:
        val = record.get(col, "")
        if val is None:
            val = ""
        row.append(str(val))

    ws.append_row(row, value_input_option="USER_ENTERED")


@st.cache_data(show_spinner=False, ttl=30)
def load_sheet_df() -> pd.DataFrame:
    ws = get_worksheet()
    records = ws.get_all_records(expected_headers=SHEET_COLUMNS)

    if not records:
        return pd.DataFrame(columns=SHEET_COLUMNS)

    df = pd.DataFrame(records)

    for c in SHEET_COLUMNS:
        if c not in df.columns:
            df[c] = ""

    return df[SHEET_COLUMNS]


# =========================
# GOOGLE DRIVE UPLOAD (RCA IMAGE)
# =========================

def upload_rca_image_to_drive(uploaded_file: Any, record_id: str) -> Dict[str, str]:
    """
    อัปโหลดไฟล์ภาพ RCA ไป Google Drive แล้วคืนค่า metadata
    หมายเหตุ: ต้อง share โฟลเดอร์ปลายทางให้ service account ก่อน
    """
    if uploaded_file is None:
        return {"file_id": "", "file_name": "", "file_url": ""}

    folder_id = str(CFG.get("GDRIVE_FOLDER_ID", "") or "").strip()
    if not folder_id:
        raise ValueError("ยังไม่ได้ตั้งค่า GDRIVE_FOLDER_ID ใน Environment Variables")

    drive = get_drive_service()

    original_name = getattr(uploaded_file, "name", "rca_image.png")
    mime_type = getattr(uploaded_file, "type", None) or "application/octet-stream"

    safe_name = f"{record_id}_{original_name}"

    file_metadata = {
        "name": safe_name,
        "parents": [folder_id],
    }

    media = MediaIoBaseUpload(
        BytesIO(uploaded_file.getvalue()),
        mimetype=mime_type,
        resumable=False,
    )

    created = drive.files().create(
        body=file_metadata,
        media_body=media,
        fields="id,name",
        supportsAllDrives=True,
    ).execute()

    file_id = created.get("id", "")
    file_name = created.get("name", safe_name)
    file_url = f"https://drive.google.com/file/d/{file_id}/view" if file_id else ""

    return {
        "file_id": file_id,
        "file_name": file_name,
        "file_url": file_url,
    }


# =========================
# DOCX EXPORT (BEFORE SAVE)
# =========================

def build_docx_report_bytes(uploaded_rca_image: Optional[Any] = None) -> bytes:
    """
    สร้างเอกสาร DOCX จากข้อมูลในฟอร์มปัจจุบัน (ก่อนบันทึก)
    """
    doc = Document()

    # Header
    doc.add_heading("รายงาน Medication Error / RCA (ก่อนบันทึก)", level=1)
    doc.add_paragraph(f"หน่วยงาน: {CFG.get('UNIT_NAME', '-')}")
    doc.add_paragraph(f"ระบบ: {CFG.get('APP_TITLE', '-')}")
    doc.add_paragraph(f"วันที่สร้างเอกสาร: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    # ข้อมูลเหตุการณ์
    doc.add_heading("1) ข้อมูลเหตุการณ์", level=2)
    t = doc.add_table(rows=0, cols=2)
    t.style = "Table Grid"

    def add_row(k: str, v: str):
        row = t.add_row().cells
        row[0].text = str(k)
        row[1].text = str(v or "")

    event_date_val = st.session_state.get("form_event_date", "")
    event_time_val = st.session_state.get("form_event_time", "")

    if isinstance(event_date_val, date):
        event_date_text = event_date_val.isoformat()
    else:
        event_date_text = str(event_date_val)

    if isinstance(event_time_val, time):
        event_time_text = event_time_val.strftime("%H:%M")
    else:
        event_time_text = str(event_time_val)

    add_row("วันที่เกิดเหตุ", event_date_text)
    add_row("เวลาเกิดเหตุ", event_time_text)
    add_row("กระบวนการที่เกิด", st.session_state.get("form_process_step", ""))
    add_row("ชื่อยา", st.session_state.get("form_drug_name", ""))
    add_row("ระดับความรุนแรง", st.session_state.get("form_severity", ""))

    # รายละเอียดเหตุการณ์
    doc.add_heading("2) รายละเอียดเหตุการณ์", level=2)
    doc.add_paragraph(st.session_state.get("form_incident_detail", "") or "-")

    # ข้อมูลเสริมในฟอร์ม
    doc.add_heading("3) ข้อมูลเสริม (จากผู้ใช้)", level=2)

    doc.add_paragraph("3.1 ไทม์ไลน์")
    doc.add_paragraph(st.session_state.get("form_timeline_text", "") or "-")

    doc.add_paragraph("3.2 การแก้ไขเบื้องต้น")
    doc.add_paragraph(st.session_state.get("form_initial_correction", "") or "-")

    doc.add_paragraph("3.3 RCA (ข้อความ)")
    doc.add_paragraph(st.session_state.get("form_rca_text", "") or "-")

    doc.add_paragraph("3.4 แผนพัฒนา")
    doc.add_paragraph(st.session_state.get("form_development_plan", "") or "-")

    # ผลวิเคราะห์ AI (ถ้ามี)
    analysis = st.session_state.get("rca_analysis_json") or {}
    plan = st.session_state.get("rca_plan_json") or {}

    if analysis:
        doc.add_heading("4) ผลวิเคราะห์ RCA จากระบบ", level=2)

        doc.add_paragraph("4.1 สรุปเหตุการณ์")
        doc.add_paragraph(str(analysis.get("event_summary", "-")))

        timeline = analysis.get("timeline", []) or []
        doc.add_paragraph("4.2 ไทม์ไลน์เหตุการณ์")
        if timeline:
            for item in timeline:
                doc.add_paragraph(f"- {item}")
        else:
            doc.add_paragraph("-")

        fishbone = analysis.get("fishbone", {}) or {}
        doc.add_paragraph("4.3 Fishbone (สรุปแบบข้อความ)")
        effect = fishbone.get("effect", "")
        if effect:
            doc.add_paragraph(f"ผลลัพธ์/เหตุการณ์: {effect}")
        for cat in (fishbone.get("categories", []) or []):
            label = str(cat.get("label", "") or "ไม่ระบุ")
            doc.add_paragraph(f"หมวด: {label}")
            for it in (cat.get("items", []) or []):
                doc.add_paragraph(f"  - {it}")

        whys = analysis.get("five_whys", []) or []
        doc.add_paragraph("4.4 5 Whys")
        if whys:
            for w in whys:
                doc.add_paragraph(f"- {w}")
        else:
            doc.add_paragraph("-")

        swiss = analysis.get("swiss_cheese", []) or []
        doc.add_paragraph("4.5 Swiss Cheese")
        if swiss:
            for row in swiss:
                line = (
                    f"[{row.get('layer','')}] "
                    f"type={row.get('type','')} | "
                    f"hole={row.get('hole','')} | "
                    f"prevention={row.get('prevention','')}"
                )
                doc.add_paragraph(f"- {line}")
        else:
            doc.add_paragraph("-")

        factors = analysis.get("contributing_factors", []) or []
        doc.add_paragraph("4.6 ปัจจัยเอื้อ/ปัจจัยร่วม")
        if factors:
            for f in factors:
                doc.add_paragraph(f"- {f}")
        else:
            doc.add_paragraph("-")

    if plan:
        doc.add_heading("5) แผนปฏิบัติการ / PDSA จากระบบ", level=2)

        pdsa = plan.get("pdsa", {}) or {}
        for key_th, key_en in [
            ("Plan", "plan"),
            ("Do", "do"),
            ("Study", "study"),
            ("Act", "act"),
        ]:
            doc.add_paragraph(f"PDSA - {key_th}")
            items = pdsa.get(key_en, []) or []
            if items:
                for it in items:
                    doc.add_paragraph(f"- {it}")
            else:
                doc.add_paragraph("-")

        ap = plan.get("action_plan", []) or []
        doc.add_paragraph("Action Plan")
        if ap:
            for i, row in enumerate(ap, 1):
                line = (
                    f"{i}) {row.get('measure','')} | "
                    f"ผู้รับผิดชอบ: {row.get('owner','')} | "
                    f"กำหนดเสร็จ: {row.get('due','')} | "
                    f"KPI: {row.get('kpi','')}"
                )
                doc.add_paragraph(line)
        else:
            doc.add_paragraph("-")

        ideas = plan.get("initiative_ideas", {}) or {}
        doc.add_paragraph("Initiative Ideas - Quick Wins (0–30 วัน)")
        for x in ideas.get("quick_wins_0_30_days", []) or []:
            doc.add_paragraph(f"- {x}")

        doc.add_paragraph("Initiative Ideas - ระยะกลาง (1–3 เดือน)")
        for x in ideas.get("mid_term_1_3_months", []) or []:
            doc.add_paragraph(f"- {x}")

        doc.add_paragraph("Initiative Ideas - ระยะยาว (3–12 เดือน)")
        for x in ideas.get("long_term_3_12_months", []) or []:
            doc.add_paragraph(f"- {x}")

        recs = plan.get("conclusion_recommendations", []) or []
        doc.add_paragraph("Conclusion & Recommendations")
        for i, x in enumerate(recs, 1):
            doc.add_paragraph(f"{i}. {x}")

        next72 = plan.get("next_72_hours", []) or []
        doc.add_paragraph("ก้าวถัดไป (ภายใน 72 ชั่วโมง)")
        for x in next72:
            doc.add_paragraph(f"- {x}")

    # แนบภาพ RCA ที่ผู้ใช้อัปโหลด (ถ้ามี)
    if uploaded_rca_image is not None:
        try:
            doc.add_heading("6) ภาพ RCA ที่แนบ", level=2)
            img_bytes = uploaded_rca_image.getvalue()
            doc.add_paragraph(f"ชื่อไฟล์: {getattr(uploaded_rca_image, 'name', '-')}")
            doc.add_picture(BytesIO(img_bytes), width=Inches(6.2))
        except Exception as e:
            doc.add_paragraph(f"(ไม่สามารถแทรกรูปลง DOCX ได้: {e})")

    out = BytesIO()
    doc.save(out)
    out.seek(0)
    return out.getvalue()


# =========================
# GEMINI / RCA ASSISTANT
# =========================

def call_gemini_json(
    prompt: str,
    api_key: str,
    image_file: Optional[Any] = None,
    timeout_sec: int = 60,
) -> Dict[str, Any]:
    """
    เรียก Gemini ผ่าน REST และบังคับ response เป็น JSON
    รองรับแนบภาพ (optional)
    """
    if not api_key:
        raise ValueError("ยังไม่ได้ตั้งค่า GEMINI_API_KEY ใน Environment Variables")

    url = (
        "https://generativelanguage.googleapis.com/v1beta/models/"
        f"gemini-2.5-flash:generateContent?key={api_key}"
    )

    parts: List[Dict[str, Any]] = [{"text": prompt}]

    if image_file is not None:
        try:
            import base64
            img_bytes = image_file.getvalue()
            mime_type = getattr(image_file, "type", None) or "image/png"
            parts.append(
                {
                    "inline_data": {
                        "mime_type": mime_type,
                        "data": base64.b64encode(img_bytes).decode("utf-8"),
                    }
                }
            )
        except Exception:
            # ถ้าอ่านรูปไม่ได้ ยังไม่ให้พังทั้ง flow
            pass

    payload = {
        "contents": [{"parts": parts}],
        "generationConfig": {"responseMimeType": "application/json"},
        "safetySettings": [
            {"category": "HARM_CATEGORY_HARASSMENT", "threshold": "BLOCK_NONE"},
            {"category": "HARM_CATEGORY_HATE_SPEECH", "threshold": "BLOCK_NONE"},
            {"category": "HARM_CATEGORY_SEXUALLY_EXPLICIT", "threshold": "BLOCK_NONE"},
            {"category": "HARM_CATEGORY_DANGEROUS_CONTENT", "threshold": "BLOCK_NONE"},
        ],
    }

    resp = requests.post(url, json=payload, timeout=timeout_sec)
    try:
        data = resp.json()
    except Exception:
        raise RuntimeError(f"Gemini API ตอบกลับไม่ใช่ JSON (HTTP {resp.status_code})")

    if not resp.ok:
        err_msg = data.get("error", {}).get("message", f"Gemini API error ({resp.status_code})")
        raise RuntimeError(err_msg)

    text = (
        data.get("candidates", [{}])[0]
        .get("content", {})
        .get("parts", [{}])[0]
        .get("text", "")
    )

    if not text:
        raise RuntimeError("Gemini ไม่ส่งผลลัพธ์กลับมา")

    # clean code fences if any
    cleaned = re.sub(r"^```(?:json)?\s*", "", text.strip(), flags=re.I)
    cleaned = re.sub(r"\s*```$", "", cleaned.strip())

    try:
        return json.loads(cleaned)
    except json.JSONDecodeError as e:
        raise RuntimeError(f"Gemini ส่ง JSON ไม่ถูกต้อง: {e}\n\nRaw response:\n{cleaned[:2000]}")


def build_analysis_prompt(incident_text: str) -> str:
    return f"""
คุณคือผู้เชี่ยวชาญด้านความปลอดภัยผู้ป่วยและ RCA ในโรงพยาบาล
โปรดวิเคราะห์เหตุการณ์ต่อไปนี้เป็นภาษาไทย และส่งกลับเป็น JSON เท่านั้น (ห้ามมี markdown ห้ามมีข้อความอื่นนอก JSON)

เหตุการณ์:
\"\"\"{incident_text}\"\"\"

โครงสร้าง JSON ที่ต้องการ:
{{
  "event_summary": "สรุปเหตุการณ์แบบกระชับ 2-4 บรรทัด",
  "timeline": [
    "เหตุการณ์ลำดับที่ 1 ...",
    "เหตุการณ์ลำดับที่ 2 ..."
  ],
  "fishbone": {{
    "effect": "เหตุการณ์/ผลลัพธ์สั้นๆ",
    "categories": [
      {{
        "label": "คน",
        "items": ["...", "..."]
      }},
      {{
        "label": "วิธีการ",
        "items": ["...", "..."]
      }}
    ]
  }},
  "five_whys": [
    "ทำไม 1: ...",
    "ทำไม 2: ...",
    "ทำไม 3: ...",
    "ทำไม 4: ...",
    "ทำไม 5: ... (รากสาเหตุ)"
  ],
  "swiss_cheese": [
    {{
      "layer": "นโยบายองค์กร",
      "type": "latent/active",
      "hole": "ช่องโหว่",
      "prevention": "ข้อเสนอป้องกัน"
    }}
  ],
  "contributing_factors": [
    "ปัจจัยเอื้อ 1",
    "ปัจจัยเอื้อ 2"
  ]
}}

ข้อกำหนด:
- fishbone.categories มีได้สูงสุด 6 หมวด
- แต่ละหมวด items สูงสุด 5 ข้อ
- swiss_cheese อย่างน้อย 4 แถว
- five_whys ให้ครบ 5 ข้อ
- ใช้ภาษาไทยล้วน
    """.strip()


def build_plan_prompt(incident_text: str, analysis_json: Dict[str, Any]) -> str:
    analysis_text = json.dumps(analysis_json, ensure_ascii=False)
    return f"""
คุณคือผู้จัดการความปลอดภัยของโรงพยาบาล
จากเหตุการณ์และผลวิเคราะห์ RCA ด้านล่าง โปรดสร้างแผนปฏิบัติการ และส่งกลับเป็น JSON เท่านั้น

เหตุการณ์:
\"\"\"{incident_text}\"\"\"

ผลวิเคราะห์:
{analysis_text}

โครงสร้าง JSON:
{{
  "pdsa": {{
    "plan": ["...","..."],
    "do": ["...","..."],
    "study": ["...","..."],
    "act": ["...","..."]
  }},
  "action_plan": [
    {{
      "measure": "มาตรการ",
      "owner": "ผู้รับผิดชอบ",
      "due": "กำหนดเสร็จ",
      "kpi": "ตัวชี้วัด",
      "risk_control": "ความเสี่ยงและแนวทางลดเสี่ยง"
    }}
  ],
  "initiative_ideas": {{
    "quick_wins_0_30_days": ["...","..."],
    "mid_term_1_3_months": ["...","..."],
    "long_term_3_12_months": ["...","..."]
  }},
  "conclusion_recommendations": [
    "ข้อเสนอแนะสำคัญข้อ 1",
    "ข้อเสนอแนะสำคัญข้อ 2",
    "ข้อเสนอแนะสำคัญข้อ 3",
    "ข้อเสนอแนะสำคัญข้อ 4",
    "ข้อเสนอแนะสำคัญข้อ 5"
  ],
  "next_72_hours": [
    "ก้าวถัดไปภายใน 72 ชั่วโมง ข้อ 1",
    "ก้าวถัดไปภายใน 72 ชั่วโมง ข้อ 2"
  ]
}}

ข้อกำหนด:
- action_plan 3-8 แถว
- recommendation ให้ 5 ข้อพอดี
- ใช้ภาษาไทย
    """.strip()


# =========================
# FISHBONE SVG (EXECUTIVE-FRIENDLY)
# =========================

def _wrap_by_chars(text: str, max_chars: int = 24, max_lines: int = 3) -> List[str]:
    s = str(text or "").strip()
    if not s:
        return []
    out: List[str] = []
    i = 0
    while i < len(s) and len(out) < max_lines:
        out.append(s[i:i + max_chars])
        i += max_chars
    if i < len(s) and out:
        # เติม … ท้ายบรรทัดสุดท้าย
        out[-1] = (out[-1][:-1] + "…") if len(out[-1]) >= 1 else "…"
    return out


def _tspans(
    lines: List[str],
    x: float,
    first_y: float,
    line_h: float = 18,
    anchor: str = "start",
    font_size: int = 13,
    font_weight: str = "400",
    fill: str = "#0f172a",
) -> str:
    if not lines:
        return ""
    chunks = []
    for i, line in enumerate(lines):
        dy = "0" if i == 0 else str(line_h)
        chunks.append(
            f'<tspan x="{x}" dy="{dy}">{html.escape(line)}</tspan>'
        )
    return (
        f'<text x="{x}" y="{first_y}" text-anchor="{anchor}" '
        f'font-size="{font_size}" font-weight="{font_weight}" '
        f'font-family="Sarabun, Noto Sans Thai, sans-serif" fill="{fill}">'
        + "".join(chunks)
        + "</text>"
    )


def fishbone_svg(effect: str, categories: List[Dict[str, Any]]) -> str:
    """
    Executive-friendly fishbone:
    - เน้นอ่านง่ายสำหรับผู้บริหาร
    - ใช้ 4 หมวดหลักบนรูป (บน 2 / ล่าง 2)
    - หมวดละ 1-2 ข้อบนรูป
    - รายละเอียดเต็มแสดงใน expander ด้านล่าง
    """
    # เตรียมหมวด
    raw = categories or []
    if not raw:
        raw = [{"label": "ยังไม่มีข้อมูล", "items": []}]

    # ใช้ 4 หมวดแรกสำหรับภาพ (ฉบับผู้บริหาร)
    raw = raw[:4]

    cats: List[Dict[str, Any]] = []
    for c in raw:
        label = str(c.get("label", "")).strip() or "ไม่ระบุ"
        items = [str(x).strip() for x in (c.get("items", []) or []) if str(x).strip()]
        cats.append({"label": label, "items": items[:2]})

    while len(cats) < 4:
        cats.append({"label": "", "items": []})

    # Canvas ขนาดใหญ่เพื่ออ่านง่ายและไม่ตกขอบ
    W, H = 2400, 1200
    spine_y = 600
    spine_x1 = 180

    head_x = 1700
    head_y = 380
    head_w = 620
    head_h = 440

    # ตำแหน่งกระดูก 4 จุด (บนซ้าย/บนขวา/ล่างซ้าย/ล่างขวา)
    anchors = [
        {"x": 820, "end_y": 280, "top": True},
        {"x": 1250, "end_y": 280, "top": True},
        {"x": 920, "end_y": 940, "top": False},
        {"x": 1350, "end_y": 940, "top": False},
    ]
    end_dx = 300

    line_layer: List[str] = []
    text_layer: List[str] = []

    for i, c in enumerate(cats):
        if not c["label"]:
            continue

        a = anchors[i]
        x = float(a["x"])
        end_y = float(a["end_y"])
        is_top = bool(a["top"])
        end_x = x - end_dx

        # เส้นกระดูกหลัก
        line_layer.append(
            f'<line x1="{x}" y1="{spine_y}" x2="{end_x}" y2="{end_y}" stroke="#334155" stroke-width="4"/>'
        )

        # เวกเตอร์สำหรับ rib
        dx = end_x - x
        dy = end_y - spine_y
        ln = (dx ** 2 + dy ** 2) ** 0.5 or 1.0
        ux, uy = dx / ln, dy / ln
        px, py = -uy, ux
        if is_top:
            px, py = -px, -py

        # กล่องหัวหมวด
        label_w = 360
        label_h = 52
        label_x = end_x - label_w - 14
        label_y = end_y - 72 if is_top else end_y + 18

        text_layer.append(
            f'<rect x="{label_x}" y="{label_y}" width="{label_w}" height="{label_h}" rx="14" '
            f'fill="#ffffff" stroke="#94a3b8" stroke-width="2"/>'
        )
        text_layer.append(
            _tspans(
                _wrap_by_chars(c["label"], max_chars=28, max_lines=1),
                x=label_x + 16,
                first_y=label_y + 33,
                line_h=18,
                anchor="start",
                font_size=17,
                font_weight="700",
            )
        )

        # Ribs + กล่องข้อความ (กันเส้นทับตัวหนังสือ)
        rib_positions = [0.38, 0.62]
        rib_len = 54

        for j, item in enumerate(c["items"][:2]):
            f = rib_positions[j]
            sx = x + dx * f
            sy = spine_y + dy * f
            ex = sx + px * rib_len
            ey = sy + py * rib_len

            line_layer.append(
                f'<line x1="{sx}" y1="{sy}" x2="{ex}" y2="{ey}" stroke="#64748b" stroke-width="3"/>'
            )

            # กล่องข้อความ rib (2 บรรทัด)
            item_lines = _wrap_by_chars(item, max_chars=34, max_lines=2)

            box_w = 410
            box_h = 56 if len(item_lines) <= 1 else 76
            box_x = ex - box_w - 10
            box_y = ey - box_h - 6 if is_top else ey + 6

            # กันหลุดซ้าย
            if box_x < 20:
                box_x = 20

            text_layer.append(
                f'<rect x="{box_x}" y="{box_y}" width="{box_w}" height="{box_h}" rx="10" '
                f'fill="#ffffff" stroke="#e2e8f0" stroke-width="1.5" opacity="0.98"/>'
            )
            text_layer.append(
                _tspans(
                    item_lines,
                    x=box_x + 12,
                    first_y=box_y + 22,
                    line_h=20,
                    anchor="start",
                    font_size=13,
                    font_weight="400",
                )
            )

    # กล่องหัวปลา (เพิ่มพื้นที่และจำนวนบรรทัด)
    effect_lines = _wrap_by_chars(effect or "เหตุการณ์ / ผลลัพธ์", max_chars=26, max_lines=8)
    effect_text = _tspans(
        effect_lines,
        x=head_x + head_w / 2,
        first_y=head_y + 98,
        line_h=28,
        anchor="middle",
        font_size=20,
        font_weight="700",
    )

    svg = f"""
    <svg viewBox="0 0 {W} {H}" width="100%" height="760" xmlns="http://www.w3.org/2000/svg">
      <defs>
        <marker id="arrowHead" markerWidth="18" markerHeight="18" refX="15" refY="9" orient="auto">
          <path d="M0,0 L18,9 L0,18 Z" fill="#0ea5e9"/>
        </marker>
      </defs>

      <!-- background -->
      <rect x="0" y="0" width="{W}" height="{H}" fill="#ffffff"/>

      <!-- spine -->
      <circle cx="{spine_x1}" cy="{spine_y}" r="12" fill="#0f172a"/>
      <line x1="{spine_x1}" y1="{spine_y}" x2="{head_x}" y2="{spine_y}"
            stroke="#0f172a" stroke-width="8" marker-end="url(#arrowHead)"/>

      <!-- lines first -->
      {''.join(line_layer)}

      <!-- head -->
      <rect x="{head_x}" y="{head_y}" width="{head_w}" height="{head_h}" rx="20"
            fill="#ffffff" stroke="#0f172a" stroke-width="4"/>
      <text x="{head_x + head_w/2}" y="{head_y + 52}" text-anchor="middle"
            font-size="22" font-weight="800"
            font-family="Sarabun, Noto Sans Thai, sans-serif" fill="#0f172a">
        เหตุการณ์ / ผลลัพธ์
      </text>

      {effect_text}

      <!-- text last -->
      {''.join(text_layer)}

      <text x="{spine_x1 - 10}" y="{spine_y - 24}" text-anchor="middle"
            font-size="14" font-weight="700"
            font-family="Sarabun, Noto Sans Thai, sans-serif" fill="#475569">สาเหตุ</text>
    </svg>
    """
    return svg


# =========================
# RENDER ANALYSIS / PLAN
# =========================

def render_analysis_result(analysis: Dict[str, Any]):
    st.subheader("🔎 ผลวิเคราะห์ RCA")

    # 1) Summary
    st.markdown("### 1) สรุปเหตุการณ์")
    st.write(analysis.get("event_summary", "-"))

    # 2) Timeline
    st.markdown("### 2) ไทม์ไลน์เหตุการณ์")
    timeline = analysis.get("timeline", []) or []
    if timeline:
        for i, item in enumerate(timeline, 1):
            st.markdown(f"- **{i}.** {item}")
    else:
        st.write("-")

    # 3) Fishbone (แสดงเฉพาะรายละเอียด ไม่แสดงภาพ)
    st.markdown("### 3) แผนผังก้างปลา (Ishikawa) — รายละเอียด")
    fishbone = analysis.get("fishbone", {}) or {}
    effect = fishbone.get("effect", "") or analysis.get("event_summary", "เหตุการณ์ / ผลลัพธ์")
    categories = fishbone.get("categories", []) or []

    st.markdown("**เหตุการณ์ / ผลลัพธ์**")
    st.write(effect if str(effect).strip() else "-")

    if categories:
        for idx, c in enumerate(categories, 1):
            label = str(c.get("label", "") or "ไม่ระบุ").strip()
            items = [str(x).strip() for x in (c.get("items", []) or []) if str(x).strip()]

            st.markdown(f"**{idx}) {label}**")
            if items:
                for item in items:
                    st.markdown(f"- {item}")
            else:
                st.markdown("- ไม่มีรายละเอียด")
    else:
        st.write("-")
    # 4) 5 Whys
    st.markdown("### 4) วิเคราะห์ทำไม-ทำไม (5 Whys)")
    whys = analysis.get("five_whys", []) or []
    if whys:
        for i, w in enumerate(whys, 1):
            st.markdown(f"{i}. {w}")
    else:
        st.write("-")

    # 5) Swiss cheese
    st.markdown("### 5) โมเดลสวิสชีส")
    swiss = analysis.get("swiss_cheese", []) or []
    if swiss:
        df_swiss = pd.DataFrame(swiss)
        df_swiss = df_swiss.rename(
            columns={
                "layer": "ชั้นระบบ",
                "type": "ประเภท",
                "hole": "รู (ช่องโหว่)",
                "prevention": "มาตรการป้องกัน",
            }
        )
        st.dataframe(df_swiss, use_container_width=True, hide_index=True)
    else:
        st.write("-")

    # 6) contributing factors
    factors = analysis.get("contributing_factors", []) or []
    if factors:
        st.markdown("### 6) ปัจจัยเอื้อ/ปัจจัยร่วม")
        for f in factors:
            st.markdown(f"- {f}")


def render_plan_result(plan: Dict[str, Any]):
    st.subheader("🎯 แผนปฏิบัติการ / PDSA")

    # PDSA table
    pdsa = plan.get("pdsa", {}) or {}
    pdsa_rows = [
        ["วางแผน (Plan)", "\n".join([f"- {x}" for x in (pdsa.get("plan", []) or [])])],
        ["ทำ (Do)", "\n".join([f"- {x}" for x in (pdsa.get("do", []) or [])])],
        ["ศึกษา (Study)", "\n".join([f"- {x}" for x in (pdsa.get("study", []) or [])])],
        ["ปรับปรุง (Act)", "\n".join([f"- {x}" for x in (pdsa.get("act", []) or [])])],
    ]
    st.markdown("### 1) PDSA")
    st.dataframe(
        pd.DataFrame(pdsa_rows, columns=["ขั้นตอน", "รายละเอียด"]),
        use_container_width=True,
        hide_index=True,
    )

    # Action plan
    st.markdown("### 2) Action Plan")
    ap = plan.get("action_plan", []) or []
    if ap:
        df_ap = pd.DataFrame(ap)
        df_ap = df_ap.rename(
            columns={
                "measure": "มาตรการ",
                "owner": "ผู้รับผิดชอบ",
                "due": "กำหนดเสร็จ",
                "kpi": "KPI(ตัวชี้วัดผลลัพธ์)",
                "risk_control": "ความเสี่ยงและแนวทางลดเสี่ยง",
            }
        )
        st.dataframe(df_ap, use_container_width=True, hide_index=True)
    else:
        st.write("-")

    # Initiative ideas
    st.markdown("### 3) Initiative Ideas")
    ideas = plan.get("initiative_ideas", {}) or {}
    col1, col2, col3 = st.columns(3)
    with col1:
        st.markdown("**Quick Wins (0–30 วัน)**")
        for x in ideas.get("quick_wins_0_30_days", []) or []:
            st.markdown(f"- {x}")
    with col2:
        st.markdown("**ระยะกลาง (1–3 เดือน)**")
        for x in ideas.get("mid_term_1_3_months", []) or []:
            st.markdown(f"- {x}")
    with col3:
        st.markdown("**ระยะยาว (3–12 เดือน)**")
        for x in ideas.get("long_term_3_12_months", []) or []:
            st.markdown(f"- {x}")

    # Conclusion & next 72h
    st.markdown("### 4) Conclusion & Recommendations")
    recs = plan.get("conclusion_recommendations", []) or []
    if recs:
        for i, x in enumerate(recs, 1):
            st.markdown(f"{i}. {x}")
    else:
        st.write("-")

    st.markdown("**ก้าวถัดไป (ภายใน 72 ชั่วโมง)**")
    next72 = plan.get("next_72_hours", []) or []
    if next72:
        for x in next72:
            st.markdown(f"- {x}")
    else:
        st.write("-")


# =========================
# FORM / SAVE
# =========================

def init_form_state_defaults():
    defaults = {
        "form_event_date": date.today(),
        "form_event_time": datetime.now().time().replace(second=0, microsecond=0),
        "form_process_step": PROCESS_OPTIONS[0],
        "form_drug_name": "",
        "form_severity": "A",
        "form_incident_detail": "",
        "form_timeline_text": "",
        "form_initial_correction": "",
        "form_rca_text": "",
        "form_development_plan": "",
        "rca_analysis_json": None,
        "rca_plan_json": None,
    }
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v


def validate_required_form() -> Tuple[bool, List[str]]:
    errs: List[str] = []
    if not st.session_state.get("form_drug_name", "").strip():
        errs.append("กรุณากรอกชื่อยา")
    if not st.session_state.get("form_incident_detail", "").strip():
        errs.append("กรุณากรอกรายละเอียดเหตุการณ์")
    return (len(errs) == 0, errs)


def create_record_from_form(
    uploaded_rca_image: Optional[Any],
    rca_image_drive_url: str = "",
) -> Dict[str, Any]:
    now = datetime.now()
    event_date_val = st.session_state.get("form_event_date")
    event_time_val = st.session_state.get("form_event_time")

    if isinstance(event_date_val, datetime):
        event_date_str = event_date_val.date().isoformat()
    elif isinstance(event_date_val, date):
        event_date_str = event_date_val.isoformat()
    else:
        event_date_str = str(event_date_val)

    if isinstance(event_time_val, datetime):
        event_time_str = event_time_val.strftime("%H:%M")
    elif isinstance(event_time_val, time):
        event_time_str = event_time_val.strftime("%H:%M")
    else:
        event_time_str = str(event_time_val)

    record = {
        "record_id": now.strftime("%Y%m%d%H%M%S%f"),
        "unit_name": CFG["UNIT_NAME"],
        "app_title": CFG["APP_TITLE"],
        "event_date": event_date_str,
        "event_time": event_time_str,
        "process_step": st.session_state.get("form_process_step", ""),
        "drug_name": st.session_state.get("form_drug_name", "").strip(),
        "severity_level": st.session_state.get("form_severity", ""),
        "incident_detail": st.session_state.get("form_incident_detail", "").strip(),
        "timeline_text": st.session_state.get("form_timeline_text", "").strip(),
        "initial_correction": st.session_state.get("form_initial_correction", "").strip(),
        "rca_text": st.session_state.get("form_rca_text", "").strip(),
        "rca_image_filename": getattr(uploaded_rca_image, "name", "") if uploaded_rca_image else "",
        "rca_image_drive_url": (rca_image_drive_url or "").strip(),
        "development_plan": st.session_state.get("form_development_plan", "").strip(),
        "created_at": now.isoformat(timespec="seconds"),
        "created_by": st.session_state.get("login_username", ""),
    }
    return record


def request_form_reset_after_save():
    """ขอให้ล้างฟอร์มในรอบถัดไป (ห้ามล้างทันทีในรอบที่กดปุ่ม)"""
    st.session_state["_reset_form_after_save"] = True
    st.session_state["_save_success_message"] = "บันทึกข้อมูลสำเร็จ ✅"


def apply_pending_form_reset():
    """ถ้ามี flag ให้ล้างฟอร์มก่อนสร้าง widget"""
    if st.session_state.get("_reset_form_after_save", False):
        st.session_state["form_drug_name"] = ""
        st.session_state["form_incident_detail"] = ""
        st.session_state["form_timeline_text"] = ""
        st.session_state["form_initial_correction"] = ""
        st.session_state["form_rca_text"] = ""
        st.session_state["form_development_plan"] = ""
        st.session_state["form_process_step"] = PROCESS_OPTIONS[0]
        st.session_state["form_severity"] = "A"
        st.session_state["form_event_date"] = date.today()
        st.session_state["form_event_time"] = datetime.now().time().replace(second=0, microsecond=0)
        st.session_state["rca_analysis_json"] = None
        st.session_state["rca_plan_json"] = None
        st.session_state["show_fishbone_preview"] = False

        # เคลียร์ file_uploader
        st.session_state.pop("form_rca_image", None)

        st.session_state["_reset_form_after_save"] = False


def render_entry_tab():
    init_form_state_defaults()
    apply_pending_form_reset()

    if st.session_state.get("_save_success_message"):
        st.success(st.session_state.pop("_save_success_message"))

    st.markdown("## 📝 บันทึกข้อมูล")

    left, right = st.columns([1.15, 1], gap="large")

    # ใช้อัปโหลดภาพ RCA ตัวเดียว ทั้งแสดงผล/ส่ง AI/ส่งขึ้น Drive
    uploaded_rca_image = None

    with left:
        st.markdown("### ข้อมูลเหตุการณ์")

        c1, c2 = st.columns(2)
        with c1:
            st.date_input("วันที่เกิดเหตุ", key="form_event_date")
        with c2:
            st.time_input("เวลาเกิดเหตุ", key="form_event_time")

        st.selectbox("กระบวนการที่เกิด", PROCESS_OPTIONS, key="form_process_step")
        st.text_input("ชื่อยา", key="form_drug_name")
        st.selectbox("ระดับความรุนแรง", SEVERITY_OPTIONS, key="form_severity")
        st.text_area("รายละเอียดเหตุการณ์", height=140, key="form_incident_detail")

        st.markdown("---")
        st.markdown("### ข้อมูลเสริม (ก่อนบันทึก)")

        st.text_area("1) ไทม์ไลน์", height=120, key="form_timeline_text")
        st.text_area("2) การแก้ไขเบื้องต้น", height=100, key="form_initial_correction")

        st.markdown("**3) RCA (ข้อความ + ภาพ)**")
        st.text_area("RCA (ข้อความ)", height=180, key="form_rca_text")
        uploaded_rca_image = st.file_uploader(
            "แนบภาพ RCA (เช่น ก้างปลา / แผนภาพ) - *จะเก็บชื่อไฟล์และลิงก์ Drive ในชีต*",
            type=["png", "jpg", "jpeg", "webp"],
            key="form_rca_image",
        )

        if uploaded_rca_image is not None:
            st.image(
                uploaded_rca_image,
                caption=f"ภาพ RCA: {uploaded_rca_image.name}",
                use_container_width=True,
            )

        st.text_area("4) แผนพัฒนา", height=140, key="form_development_plan")

        st.markdown("---")

        # ปุ่มดาวน์โหลด DOCX ก่อนบันทึก
        try:
            docx_bytes = build_docx_report_bytes(uploaded_rca_image=uploaded_rca_image)
            st.download_button(
                "📄 ดาวน์โหลดรายงาน DOCX (ก่อนบันทึก)",
                data=docx_bytes,
                file_name=f"RCA_Report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
        except Exception as e:
            st.caption(f"ยังไม่สามารถสร้าง DOCX ได้: {e}")

        # ปุ่มบันทึก (พร้อมอัปโหลดภาพไป Drive ถ้ามี)
        if st.button("💾 บันทึกข้อมูล", type="primary", use_container_width=True):
            ok, errs = validate_required_form()
            if not ok:
                for e in errs:
                    st.error(e)
            else:
                try:
                    # 1) สร้าง record จากฟอร์มก่อน เพื่อได้ record_id
                    record = create_record_from_form(uploaded_rca_image=uploaded_rca_image)

                    # 2) ถ้ามีภาพ → อัปโหลดขึ้น Google Drive แล้วใส่ชื่อไฟล์/ลิงก์กลับเข้า record
                    if uploaded_rca_image is not None:
                        drive_meta = upload_rca_image_to_drive(
                            uploaded_rca_image,
                            record_id=record["record_id"],
                        )
                        record["rca_image_filename"] = drive_meta.get("file_name", "") or getattr(uploaded_rca_image, "name", "")
                        record["rca_image_drive_url"] = drive_meta.get("file_url", "") or ""

                    # 3) บันทึกลง Google Sheets
                    append_record_to_sheet(record)

                    # 4) refresh cache ประวัติ (ถ้าใช้ cache_data)
                    try:
                        load_sheet_df.clear()
                    except Exception:
                        pass

                    # 5) ขอ reset ฟอร์มใน run ถัดไป แล้ว rerun
                    request_form_reset_after_save()
                    st.rerun()

                except Exception as e:
                    st.exception(e)

    with right:
        st.markdown("### 🧸 RCA Assistant")
        st.caption("ระบบจะวิเคราะห์จากรายละเอียดเหตุการณ์ แล้วแสดงผลให้ตรวจทาน จากนั้นคัดลอกไปกรอกในฟอร์มเองก่อนบันทึก")

        st.info(
            "หลักการใช้งาน: ปุ่ม RCA Assistant จะ **ไม่บันทึกลง Google Sheets** โดยอัตโนมัติ\n"
            "→ ผู้ใช้ตรวจทานผลลัพธ์ ก่อนนำไปกรอกไฟล์เอง แล้วค่อยกด **บันทึกข้อมูล**"
        )

        # ปุ่ม AI
        if st.button("🧸 RCA Assistant", use_container_width=True):
            incident_text = st.session_state.get("form_incident_detail", "").strip()
            if not incident_text:
                st.warning("กรุณากรอกรายละเอียดเหตุการณ์ก่อน")
            else:
                try:
                    with st.spinner("กำลังวิเคราะห์ RCA..."):
                        analysis = call_gemini_json(
                            prompt=build_analysis_prompt(incident_text),
                            api_key=CFG["GEMINI_API_KEY"],
                            image_file=uploaded_rca_image,
                            timeout_sec=90,
                        )
                        plan = call_gemini_json(
                            prompt=build_plan_prompt(incident_text, analysis),
                            api_key=CFG["GEMINI_API_KEY"],
                            timeout_sec=90,
                        )

                        st.session_state.rca_analysis_json = analysis
                        st.session_state.rca_plan_json = plan

                    st.success("วิเคราะห์เสร็จแล้ว ✅")
                except Exception as e:
                    st.error(f"RCA Assistant error: {e}")

        # แสดงผล AI ถ้ามี
        analysis = st.session_state.get("rca_analysis_json")
        plan = st.session_state.get("rca_plan_json")

        if analysis:
            render_analysis_result(analysis)

        if plan:
            st.markdown("---")
            render_plan_result(plan)


# =========================
# HISTORY TAB (with date fixes)
# =========================

def parse_event_datetime_columns(df: pd.DataFrame) -> pd.DataFrame:
    """
    แก้ปัญหา date:
    - NaT ใน date_input
    - dtype datetime64[ns] เทียบกับ date ไม่ได้
    """
    out = df.copy()

    out["event_date"] = out.get("event_date", "").astype(str).str.strip()
    out["event_time"] = out.get("event_time", "").astype(str).str.strip()

    out["_event_date_dt"] = pd.to_datetime(out["event_date"], errors="coerce")

    out["_event_datetime"] = pd.to_datetime(
        out["event_date"].astype(str) + " " + out["event_time"].astype(str),
        errors="coerce",
    )

    out["_event_date_only"] = out["_event_date_dt"].dt.date
    return out


def render_history_tab():
    st.markdown("## 📚 ดูข้อมูลย้อนหลัง")

    try:
        df = load_sheet_df()
    except Exception as e:
        st.error(f"โหลดข้อมูลจาก Google Sheets ไม่สำเร็จ: {e}")
        return

    if df.empty:
        st.info("ยังไม่มีข้อมูลใน Google Sheets")
        return

    df = parse_event_datetime_columns(df)

    valid_dates_series = df["_event_date_dt"].dropna()
    if valid_dates_series.empty:
        min_d = date.today()
        max_d = date.today()
    else:
        min_d = valid_dates_series.min().date()
        max_d = valid_dates_series.max().date()

    if max_d < min_d:
        min_d, max_d = max_d, min_d

    # Filters
    st.markdown("### ตัวกรอง")
    c1, c2, c3, c4 = st.columns([1, 1, 1, 1.4])

    with c1:
        start_date = st.date_input("วันที่เริ่ม", value=min_d, key="hist_start")
    with c2:
        end_date = st.date_input("วันที่สิ้นสุด", value=max_d, key="hist_end")
    with c3:
        sev_selected = st.multiselect(
            "ระดับความรุนแรง",
            options=sorted([x for x in df["severity_level"].dropna().astype(str).unique() if x]),
            default=[],
            key="hist_sev",
        )
    with c4:
        keyword = st.text_input("ค้นหา (ชื่อยา/รายละเอียด)", key="hist_kw").strip()

    proc_selected = st.multiselect(
        "กระบวนการที่เกิด",
        options=sorted([x for x in df["process_step"].dropna().astype(str).unique() if x]),
        default=[],
        key="hist_proc",
    )

    if start_date > end_date:
        st.warning("วันที่เริ่มมากกว่าวันที่สิ้นสุด ระบบจะสลับให้โดยอัตโนมัติ")
        start_date, end_date = end_date, start_date

    m = pd.Series(True, index=df.index)

    m &= df["_event_date_only"].notna()
    m &= (df["_event_date_only"] >= start_date) & (df["_event_date_only"] <= end_date)

    if sev_selected:
        m &= df["severity_level"].astype(str).isin(sev_selected)

    if proc_selected:
        m &= df["process_step"].astype(str).isin(proc_selected)

    if keyword:
        kw = keyword.lower()
        m &= (
            df["drug_name"].astype(str).str.lower().str.contains(kw, na=False)
            | df["incident_detail"].astype(str).str.lower().str.contains(kw, na=False)
            | df["rca_text"].astype(str).str.lower().str.contains(kw, na=False)
            | df["development_plan"].astype(str).str.lower().str.contains(kw, na=False)
        )

    filtered = df[m].copy()

    filtered["_created_at_dt"] = pd.to_datetime(filtered.get("created_at", ""), errors="coerce")
    filtered = filtered.sort_values(
        by=["_event_datetime", "_created_at_dt"],
        ascending=False,
        na_position="last",
    )

    st.markdown(f"**ผลลัพธ์ทั้งหมด:** {len(filtered):,} รายการ")

    if not filtered.empty:
        s1, s2, s3 = st.columns(3)
        with s1:
            st.metric("จำนวนรายการ", f"{len(filtered):,}")
        with s2:
            st.metric(
                "จำนวนยาไม่ซ้ำ",
                f"{filtered['drug_name'].astype(str).replace('', pd.NA).dropna().nunique():,}",
            )
        with s3:
            st.metric(
                "หน่วยงาน",
                str(filtered["unit_name"].astype(str).replace('', pd.NA).dropna().nunique()),
            )

    display_cols = [
        "event_date", "event_time", "process_step", "drug_name", "severity_level",
        "incident_detail", "timeline_text", "initial_correction", "rca_text",
        "rca_image_filename", "rca_image_drive_url", "development_plan", "created_at", "created_by"
    ]

    for c in display_cols:
        if c not in filtered.columns:
            filtered[c] = ""

    st.dataframe(
        filtered[display_cols],
        use_container_width=True,
        hide_index=True,
        column_config={
            "event_date": "วันที่",
            "event_time": "เวลา",
            "process_step": "กระบวนการ",
            "drug_name": "ชื่อยา",
            "severity_level": "ระดับ",
            "incident_detail": "รายละเอียดเหตุการณ์",
            "timeline_text": "ไทม์ไลน์",
            "initial_correction": "การแก้ไขเบื้องต้น",
            "rca_text": "RCA (ข้อความ)",
            "rca_image_filename": "ไฟล์ภาพ RCA",
            "rca_image_drive_url": "ลิงก์ภาพ RCA (Drive)",
            "development_plan": "แผนพัฒนา",
            "created_at": "เวลาบันทึก",
            "created_by": "ผู้บันทึก",
        },
    )

    # download csv
    csv_bytes = filtered[display_cols].to_csv(index=False).encode("utf-8-sig")
    st.download_button(
        "⬇️ ดาวน์โหลดผลลัพธ์ (CSV)",
        data=csv_bytes,
        file_name=f"med_error_history_{CFG['UNIT_NAME']}_{datetime.now().strftime('%Y%m%d_%H%M')}.csv",
        mime="text/csv",
        use_container_width=False,
    )

    # detail viewer
    with st.expander("🔍 ดูรายละเอียดรายรายการ (เลือกจากตารางด้านล่างสุด 20 รายการ)"):
        preview = filtered.head(20).copy()
        if preview.empty:
            st.write("ไม่มีข้อมูล")
        else:
            labels = []
            for _, r in preview.iterrows():
                labels.append(
                    f"{r.get('event_date','')} {r.get('event_time','')} | {r.get('drug_name','-')} | ระดับ {r.get('severity_level','-')}"
                )
            selected_idx = st.selectbox(
                "เลือกเหตุการณ์",
                options=list(range(len(labels))),
                format_func=lambda i: labels[i],
            )
            row = preview.iloc[int(selected_idx)]

            st.markdown("### รายละเอียดเหตุการณ์")
            st.write(row.get("incident_detail", ""))

            st.markdown("### ไทม์ไลน์")
            st.write(row.get("timeline_text", ""))

            st.markdown("### การแก้ไขเบื้องต้น")
            st.write(row.get("initial_correction", ""))

            st.markdown("### RCA")
            st.write(row.get("rca_text", ""))

            drive_url = str(row.get("rca_image_drive_url", "")).strip()
            if drive_url:
                st.markdown("### ลิงก์ภาพ RCA (Google Drive)")
                st.markdown(f"[เปิดไฟล์ภาพ RCA บน Google Drive]({drive_url})")

            st.markdown("### แผนพัฒนา")
            st.write(row.get("development_plan", ""))

            if str(row.get("rca_image_filename", "")).strip():
                st.caption(f"แนบภาพไว้ตอนบันทึก: {row.get('rca_image_filename')}")


# =========================
# MAIN
# =========================

def render_header():
    st.markdown(f"# 🏡 {CFG['APP_TITLE']}")
    st.caption(f"หน่วยงาน: {CFG['UNIT_NAME']}  |  บันทึกอุบัติการณ์ในสถานพยาบาลปฐมภูมิ")

    c1, c2 = st.columns([1, 6])
    with c1:
        if st.button("🚪 Logout"):
            st.session_state.authenticated = False
            st.session_state.login_username = ""
            st.session_state.show_fishbone_preview = False
            st.rerun()


def check_required_env():
    missing = []
    for key in ["GSHEET_URL", "GCP_SERVICE_ACCOUNT_JSON"]:
        if not CFG.get(key):
            missing.append(key)

    if missing:
        st.error("ยังตั้งค่า Environment Variables ไม่ครบ: " + ", ".join(missing))
        st.stop()

    # แจ้งเตือนแบบไม่บล็อก ถ้ายังไม่ได้ตั้งค่าโฟลเดอร์ Drive
    if not str(CFG.get("GDRIVE_FOLDER_ID", "") or "").strip():
        st.warning(
            "ยังไม่ได้ตั้งค่า GDRIVE_FOLDER_ID → หากแนบภาพ RCA แล้วกดบันทึก ระบบจะอัปโหลดภาพไป Google Drive ไม่ได้"
        )


def main():
    ensure_auth_state()

    if not st.session_state.authenticated:
        render_login()
        return

    check_required_env()

    render_header()
    st.markdown("---")

    tab1, tab2 = st.tabs(["บันทึกข้อมูล", "ดูข้อมูลย้อนหลัง"])

    with tab1:
        render_entry_tab()

    with tab2:
        render_history_tab()


if __name__ == "__main__":
    main()
